from __future__ import annotations
import json
import os
import hashlib
import re
import tempfile
from collections.abc import Callable
from dataclasses import asdict, dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any

from library.parse_options import DocumentParseOptions
from library.pdf_ocr_pipeline import (
    CHECKPOINT_NAMESPACE_FILENAME,
    OCRPipelineOptions,
    PADDLE_OCR_LOCK,
    _count_completed_chunks,
    _count_fully_completed_chunks,
    _count_pdf_chunks,
    _first_missing_page,
    _get_pdf_page_batch_size,
    _get_pdf_page_count,
    _get_pdf_ocr_checkpoint_root,
    _get_paddle_ocr_settings,
    _render_pdf_page_range_resilient,
    release_paddle_ocr_resources,
    run_pdf_ocr_pipeline,
)
from library.ocr_cleaner import maybe_clean_parsed_document
from settings import PROJECT_ROOT


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
TEXT_CHAR_PATTERN = re.compile(r"[\u4e00-\u9fffA-Za-z0-9]")
FORMULA_IMAGE_REF_PATTERN = re.compile(r"img_in_formula_box_[^\s\"')>]+", re.IGNORECASE)
PDF_IMAGE_MARKDOWN_PATTERN = re.compile(r"!\[[^\]]*\]\((?:data:[^)]*|imgs/[^)\s]+)\)")
PDF_IMAGE_HTML_PATTERN = re.compile(r"<img[^>]*?src=[\"'](?:data:[^\"']+|imgs/[^\"']+)[\"'][^>]*>", re.IGNORECASE)
PDF_IMAGE_PATH_PATTERN = re.compile(r"\bimgs/[^\s)\"']+")
PDF_TEXT_CHAR_THRESHOLD = 24
PDF_OCR_PREVIEW_MAX_PAGES = 2
PADDLEOCR_VL15_RUNTIME_ENV = "PADDLEOCR_VL15_RUNTIME"
PADDLEOCR_VL15_DEVICE_ENV = "PADDLEOCR_VL15_DEVICE"
PADDLEOCR_VL15_CACHE_HOME_ENV = "PADDLE_PDX_CACHE_HOME"
PADDLEOCR_VL15_MODEL_SOURCE_ENV = "PADDLE_PDX_MODEL_SOURCE"
PADDLEOCR_VL15_DISABLE_MODEL_SOURCE_CHECK_ENV = "PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"
PADDLEOCR_VL15_DISABLE_DEVICE_FALLBACK_ENV = "PADDLE_PDX_DISABLE_DEVICE_FALLBACK"
PADDLEOCR_VL15_PROVIDER = "paddleocr_vl_1_5/local"
ParseProgressCallback = Callable[[str, int, dict[str, object] | None], None]


@dataclass
class ParsedTable:
    page: int | None = None
    markdown: str = ""
    html: str = ""


@dataclass
class ParsedBlock:
    page_number: int
    block_id: str
    text: str = ""
    bbox: list[float] = field(default_factory=list)
    score: float | None = None
    block_type: str = "text"
    latex: str | None = None


@dataclass
class ParsedPage:
    page_number: int
    width: float = 0.0
    height: float = 0.0
    text: str = ""
    markdown: str = ""
    blocks: list[ParsedBlock] = field(default_factory=list)


@dataclass
class ParsedDocument:
    text: str
    markdown: str
    provider: str
    used_ocr: bool = False
    pages: list[ParsedPage] = field(default_factory=list)
    tables: list[ParsedTable] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    cleanup_report: dict[str, Any] = field(default_factory=dict)
    cleanup_score: float | None = None
    raw_text: str = ""
    raw_markdown: str = ""
    markdown_image_roots: list[str] = field(default_factory=list)
    markdown_images: dict[str, Any] = field(default_factory=dict, repr=False)

    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "markdown": self.markdown,
            "provider": self.provider,
            "used_ocr": self.used_ocr,
            "pages": [
                {
                    "page_number": page.page_number,
                    "width": page.width,
                    "height": page.height,
                    "text": page.text,
                    "markdown": page.markdown,
                    "blocks": [
                        {
                            "page_number": block.page_number,
                            "block_id": block.block_id,
                            "text": block.text,
                            "bbox": list(block.bbox),
                            "score": block.score,
                            "block_type": block.block_type,
                            "latex": block.latex,
                        }
                        for block in page.blocks
                    ],
                }
                for page in self.pages
            ],
            "tables": [
                {
                    "page": table.page,
                    "markdown": table.markdown,
                    "html": table.html,
                }
                for table in self.tables
            ],
            "warnings": list(self.warnings),
            "cleanup_report": dict(self.cleanup_report or {}),
            "cleanup_score": self.cleanup_score,
            "raw_text": self.raw_text,
            "raw_markdown": self.raw_markdown,
            "markdown_image_roots": list(self.markdown_image_roots or []),
        }


def parse_bytes(
    data: bytes,
    filename: str,
    mime: str,
    options: DocumentParseOptions | None = None,
) -> str:
    options = options or DocumentParseOptions()
    parsed = parse_document(data, filename, mime, options=options)
    return options.select_output(text=parsed.text, markdown=parsed.markdown)


def parse_document(
    data: bytes,
    filename: str,
    mime: str,
    options: DocumentParseOptions | None = None,
    progress_callback: ParseProgressCallback | None = None,
    cache_namespace: str | None = None,
) -> ParsedDocument:
    try:
        options = options or DocumentParseOptions()
        suffix = Path(filename).suffix.lower()

        if options.should_use_pdf_ocr(filename, mime):
            return _finalize_pdf_parsed_document(
                _parse_pdf_with_options(
                    data,
                    filename,
                    options,
                    progress_callback=progress_callback,
                    cache_namespace=cache_namespace,
                ),
                options,
            )
        if suffix == ".pdf" or mime == "application/pdf":
            return _finalize_pdf_parsed_document(
                _parse_pdf(data, filename, progress_callback=progress_callback, cache_namespace=cache_namespace),
                options,
            )
        if suffix in {".docx", ".doc"}:
            return _finalize_parsed_document(_parse_docx(data), options)
        if suffix in IMAGE_SUFFIXES or mime.startswith("image/"):
            return _finalize_parsed_document(_parse_image(data, suffix or ".png", options=options), options)
        return _finalize_parsed_document(_parse_text_document(data), options)
    finally:
        release_paddle_parser_resources()


def parse_preview_document(
    data: bytes,
    filename: str,
    mime: str,
    options: DocumentParseOptions | None = None,
) -> ParsedDocument:
    try:
        options = options or DocumentParseOptions()
        suffix = Path(filename).suffix.lower()
        is_pdf = suffix == ".pdf" or mime == "application/pdf"
        if is_pdf:
            return _finalize_pdf_parsed_document(
                _parse_pdf_with_options(data, filename, options, max_pages=PDF_OCR_PREVIEW_MAX_PAGES),
                options,
            )

        return _finalize_parsed_document(_parse_document_inner(data, filename, mime, options=options), options)
    finally:
        release_paddle_parser_resources()


def _finalize_parsed_document(document: ParsedDocument, options: DocumentParseOptions) -> ParsedDocument:
    return maybe_clean_parsed_document(document, raw_ocr_mode=options.use_raw_ocr_mode())


def _finalize_pdf_parsed_document(document: ParsedDocument, options: DocumentParseOptions) -> ParsedDocument:
    if not options.should_preserve_pdf_image_content():
        document = _strip_pdf_embedded_image_content(document)
    return _finalize_parsed_document(document, options)


def _parse_document_inner(
    data: bytes,
    filename: str,
    mime: str,
    options: DocumentParseOptions | None = None,
    progress_callback: ParseProgressCallback | None = None,
    cache_namespace: str | None = None,
) -> ParsedDocument:
    options = options or DocumentParseOptions()
    suffix = Path(filename).suffix.lower()

    if options.should_use_pdf_ocr(filename, mime):
        return _parse_pdf_with_options(
            data,
            filename,
            options,
            progress_callback=progress_callback,
            cache_namespace=cache_namespace,
        )
    if suffix == ".pdf" or mime == "application/pdf":
        return _parse_pdf(data, filename, progress_callback=progress_callback, cache_namespace=cache_namespace)
    if suffix in {".docx", ".doc"}:
        return _parse_docx(data)
    if suffix in IMAGE_SUFFIXES or mime.startswith("image/"):
        return _parse_image(data, suffix or ".png", options=options)
    return _parse_text_document(data)


def serialize_parsed_document(document: ParsedDocument) -> str:
    return json.dumps(document.to_dict(), ensure_ascii=False)


def deserialize_parsed_document(payload: str) -> ParsedDocument | None:
    try:
        raw = json.loads(payload)
    except (TypeError, ValueError):
        return None

    if not isinstance(raw, dict):
        return None

    pages = [
        ParsedPage(
            page_number=int(item.get("page_number", index + 1)),
            width=float(item.get("width") or 0.0),
            height=float(item.get("height") or 0.0),
            text=str(item.get("text") or ""),
            markdown=str(item.get("markdown") or ""),
            blocks=[
                ParsedBlock(
                    page_number=int(block.get("page_number") or item.get("page_number") or index + 1),
                    block_id=str(block.get("block_id") or f"p{index + 1}-b{block_index}"),
                    text=str(block.get("text") or ""),
                    bbox=[
                        float(value)
                        for value in (block.get("bbox") or [])
                        if isinstance(value, (int, float))
                    ],
                    score=float(block.get("score")) if isinstance(block.get("score"), (int, float)) else None,
                    block_type=str(block.get("block_type") or "text"),
                    latex=str(block.get("latex")) if block.get("latex") is not None else None,
                )
                for block_index, block in enumerate(item.get("blocks") or [], start=1)
                if isinstance(block, dict)
            ],
        )
        for index, item in enumerate(raw.get("pages") or [])
        if isinstance(item, dict)
    ]
    tables = [
        ParsedTable(
            page=item.get("page"),
            markdown=str(item.get("markdown") or ""),
            html=str(item.get("html") or ""),
        )
        for item in raw.get("tables") or []
        if isinstance(item, dict)
    ]
    return ParsedDocument(
        text=str(raw.get("text") or ""),
        markdown=str(raw.get("markdown") or ""),
        provider=str(raw.get("provider") or "unknown"),
        used_ocr=bool(raw.get("used_ocr")),
        pages=pages,
        tables=tables,
        warnings=[str(item) for item in raw.get("warnings") or [] if str(item).strip()],
        cleanup_report=dict(raw.get("cleanup_report") or {}),
        cleanup_score=float(raw.get("cleanup_score")) if isinstance(raw.get("cleanup_score"), (int, float)) else None,
        raw_text=str(raw.get("raw_text") or raw.get("text") or ""),
        raw_markdown=str(raw.get("raw_markdown") or raw.get("markdown") or raw.get("text") or ""),
        markdown_image_roots=[str(item) for item in raw.get("markdown_image_roots") or [] if str(item).strip()],
    )


def _parse_pdf_with_options(
    data: bytes,
    filename: str,
    options: DocumentParseOptions,
    max_pages: int | None = None,
    progress_callback: ParseProgressCallback | None = None,
    cache_namespace: str | None = None,
) -> ParsedDocument:
    if options.should_use_vl15_pipeline():
        return _parse_pdf_with_paddleocr_vl15(
            data,
            filename,
            max_pages=max_pages,
            progress_callback=progress_callback,
        )
    with PADDLE_OCR_LOCK:
        pipeline_options = options.to_pipeline_options()
        pipeline_options.cache_namespace = cache_namespace
        if max_pages is not None:
            pipeline_options.max_pages = max_pages
        layout_result: ParsedDocument | None = None
        layout_score = 0.0
        if options.should_use_layout_pipeline():
            layout_result = _parse_pdf_with_pp_structure(
                data,
                filename,
                pipeline_options,
                progress_callback,
            )
            if layout_result is not None and _document_has_meaningful_text(layout_result):
                layout_result.warnings.insert(0, f"parse_preset={options.preset}")
                if _document_has_unresolved_formula_images(layout_result):
                    layout_result.warnings.append(
                        "PP-StructureV3 left formula images unresolved; trying PP-OCRv5 fallback."
                    )
                else:
                    layout_score = _document_quality_score(layout_result)
                    if layout_score >= 6.5:
                        return layout_result
                    layout_result.warnings.append(
                        f"PP-StructureV3 quality score {layout_score:.2f} below gate; trying PP-OCRv5 fallback."
                    )

        result = run_pdf_ocr_pipeline(
            data,
            filename,
            options=pipeline_options,
            progress_callback=_ocr_progress_callback(
                progress_callback,
                stage="ocr_fallback" if layout_result is not None else "ocr",
                start=70 if layout_result is not None else 15,
                end=75 if layout_result is not None else 70,
            ),
        )
        ocr_document = _document_from_ocr_result(result, warnings_prefix=[f"parse_preset={options.preset}"])
        if layout_result is not None and _document_has_meaningful_text(layout_result):
            if _document_has_unresolved_formula_images(layout_result):
                ocr_document.warnings.insert(
                    1,
                    "PP-OCRv5 fallback selected because PP-StructureV3 left formula images unresolved.",
                )
                return ocr_document
            ocr_score = _document_quality_score(ocr_document)
            if ocr_score >= layout_score:
                ocr_document.warnings.insert(
                    1,
                    f"PP-OCRv5 fallback selected over PP-StructureV3 ({ocr_score:.2f} >= {layout_score:.2f}).",
                )
                return ocr_document
            layout_result.warnings.append(
                f"PP-StructureV3 selected over PP-OCRv5 fallback ({layout_score:.2f} > {ocr_score:.2f})."
            )
            return layout_result
        return ocr_document


def _parse_pdf_with_paddleocr_vl15(
    data: bytes,
    filename: str,
    *,
    max_pages: int | None = None,
    progress_callback: ParseProgressCallback | None = None,
) -> ParsedDocument:
    settings = _get_paddleocr_vl15_runtime_settings()
    return _parse_pdf_with_local_paddleocr_vl15(
        data,
        filename,
        max_pages=max_pages,
        progress_callback=progress_callback,
        settings=settings,
    )


def _parse_pdf_with_local_paddleocr_vl15(
    data: bytes,
    filename: str,
    *,
    max_pages: int | None = None,
    progress_callback: ParseProgressCallback | None = None,
    settings: dict[str, Any],
) -> ParsedDocument:
    pipeline = _get_local_paddleocr_vl15_pipeline(settings)

    total_page_count = _safe_pdf_page_count(data)
    request_bytes = data
    target_page_count = total_page_count
    warnings: list[str] = []
    if max_pages is not None:
        request_bytes, total_page_count, target_page_count = _slice_pdf_bytes(data, max_pages=max_pages)
        if total_page_count is not None and target_page_count is not None and total_page_count > target_page_count:
            warnings.append(f"VL1.5 preview limited to first {target_page_count} of {total_page_count} pages.")

    effective_page_count = target_page_count or total_page_count or 0
    if progress_callback is not None:
        progress_callback(
            "vl15",
            15,
            {
                "status": "initializing",
                "done_pages": 0,
                "total_pages": effective_page_count,
                "device": settings["device"],
            },
        )

    page_results: list[Any] = []
    restructure_result: Any | None = None
    restructure_warning = ""
    tmp_path: str | None = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=Path(filename).suffix or ".pdf")
        with os.fdopen(fd, "wb") as tmp:
            tmp.write(request_bytes)
            tmp.flush()

        with PADDLE_OCR_LOCK:
            for index, page_result in enumerate(
                pipeline.predict(
                    input=tmp_path,
                    use_doc_orientation_classify=False,
                    use_doc_unwarping=False,
                    use_layout_detection=True,
                    use_chart_recognition=False,
                    use_seal_recognition=False,
                    format_block_content=False,
                    merge_layout_blocks=True,
                ),
                start=1,
            ):
                page_results.append(page_result)
                if progress_callback is not None:
                    progress_callback(
                        "vl15",
                        min(62, 15 + int((index / max(1, effective_page_count or index)) * 47)),
                        {
                            "status": "page_parsed",
                            "done_pages": index,
                            "total_pages": effective_page_count or index,
                            "device": settings["device"],
                        },
                    )

            if page_results:
                merged_results = list(
                    pipeline.restructure_pages(
                        page_results,
                        merge_tables=True,
                        relevel_titles=True,
                        concatenate_pages=True,
                    )
                )
                if merged_results:
                    restructure_result = merged_results[0]
    except Exception as exc:
        raise RuntimeError(f"本地 PaddleOCR-VL1.5 GPU 解析失败：{exc}") from exc
    finally:
        if tmp_path:
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    if not page_results:
        return ParsedDocument(
            text="",
            markdown="",
            provider=PADDLEOCR_VL15_PROVIDER,
            used_ocr=True,
            warnings=["parse_preset=vl15", *warnings, "PaddleOCR-VL1.5 returned no page results."],
        )

    if progress_callback is not None:
        progress_callback(
            "vl15",
            72,
            {
                "status": "restructured" if restructure_result is not None else "layout_ready",
                "done_pages": len(page_results),
                "total_pages": len(page_results),
                "device": settings["device"],
            },
        )

    document = _document_from_local_vl15_results(
        page_results,
        restructure_result,
        source_bytes=request_bytes,
        filename=filename,
        warnings=warnings,
    )
    if restructure_warning:
        document.warnings.append(restructure_warning)
    document.warnings.insert(1, f"vl15_device={settings['device']}")
    document.warnings.insert(0, "parse_preset=vl15")
    return document


def _document_from_local_vl15_results(
    page_results: list[Any],
    restructure_result: Any | None,
    *,
    source_bytes: bytes,
    filename: str,
    warnings: list[str] | None = None,
) -> ParsedDocument:
    pages: list[ParsedPage] = []
    markdown_images: dict[str, Any] = {}

    for fallback_page_number, page_result in enumerate(page_results, start=1):
        page_number = _local_vl15_page_number(page_result, fallback_page_number)
        markdown_bundle = _local_vl15_markdown_bundle(page_result)
        page_markdown = str(markdown_bundle.get("markdown_texts") or "").strip()
        page_markdown_images = _extract_markdown_image_payload(markdown_bundle.get("markdown_images"))
        if page_markdown_images:
            prefixed_images, path_mapping = _prefix_markdown_image_paths(page_markdown_images, page_number=page_number)
            markdown_images.update(prefixed_images)
            page_markdown = _replace_markdown_image_paths(page_markdown, path_mapping)
        else:
            page_markdown = _strip_pdf_image_markup(page_markdown)
        page_blocks = _local_vl15_blocks(page_result, page_number)
        page_text = _markdown_to_text(page_markdown).strip()
        if not page_text:
            page_text = "\n".join(block.text for block in page_blocks if block.text.strip()).strip()
        if not page_blocks and page_text:
            page_blocks = _synthetic_blocks_from_text(page_text, page_number)
        pages.append(
            ParsedPage(
                page_number=page_number,
                width=float(page_result.get("width") or 0.0),
                height=float(page_result.get("height") or 0.0),
                text=page_text,
                markdown=page_markdown or page_text,
                blocks=page_blocks,
            )
        )

    merged_markdown = ""
    merged_text = ""
    if restructure_result is not None:
        merged_markdown_bundle = _local_vl15_markdown_bundle(restructure_result)
        merged_markdown = str(merged_markdown_bundle.get("markdown_texts") or "").strip()
        merged_markdown_images = _extract_markdown_image_payload(merged_markdown_bundle.get("markdown_images"))
        if merged_markdown_images:
            prefixed_images, path_mapping = _prefix_markdown_image_paths(merged_markdown_images, page_number=0)
            markdown_images.update(prefixed_images)
            merged_markdown = _replace_markdown_image_paths(merged_markdown, path_mapping)
        else:
            merged_markdown = _strip_pdf_image_markup(merged_markdown)
        merged_text = _markdown_to_text(merged_markdown).strip()

    if not merged_markdown:
        merged_markdown = "\n\n".join(page.markdown for page in pages if page.markdown.strip())
    if not merged_text:
        merged_text = _markdown_to_text(merged_markdown).strip()
    if not merged_text:
        merged_text = "\n\n".join(page.text for page in pages if page.text.strip())

    document = ParsedDocument(
        text=merged_text.strip(),
        markdown=(merged_markdown.strip() or merged_text.strip()),
        provider=PADDLEOCR_VL15_PROVIDER,
        used_ocr=True,
        pages=pages,
        warnings=list(warnings or []),
        markdown_images=markdown_images,
    )
    if markdown_images:
        assets_root = _build_vl15_assets_root(source_bytes, filename)
        roots = _persist_markdown_images(markdown_images, assets_root)
        document.markdown_image_roots = roots
        document.markdown_images = {}
    return document


def _local_vl15_page_number(result: Any, fallback_page_number: int) -> int:
    raw_page_index = result.get("page_index")
    if isinstance(raw_page_index, int) and raw_page_index >= 0:
        return raw_page_index + 1
    return fallback_page_number


def _local_vl15_markdown_bundle(result: Any) -> dict[str, Any]:
    payload = getattr(result, "markdown", None)
    return payload if isinstance(payload, dict) else {}


def _local_vl15_blocks(result: Any, page_number: int) -> list[ParsedBlock]:
    parsed_blocks: list[ParsedBlock] = []
    for index, block in enumerate(result.get("parsing_res_list") or [], start=1):
        text = _strip_pdf_image_markup(str(getattr(block, "content", "") or "")).strip()
        block_type = str(getattr(block, "label", "") or "text")
        bbox = [
            float(value)
            for value in (getattr(block, "bbox", None) or [])
            if isinstance(value, (int, float))
        ]
        parsed_blocks.append(
            ParsedBlock(
                page_number=page_number,
                block_id=f"p{page_number}-vl15-{index}",
                text=text,
                bbox=bbox if len(bbox) == 4 else [],
                block_type=block_type,
            )
        )
    return parsed_blocks


def _get_local_paddleocr_vl15_pipeline(settings: dict[str, Any]):
    pipeline = getattr(_get_local_paddleocr_vl15_pipeline, "_pipeline", None)
    cached_device = getattr(_get_local_paddleocr_vl15_pipeline, "_device", None)
    if pipeline is not None and cached_device == settings["device"]:
        return pipeline

    with PADDLE_OCR_LOCK:
        pipeline = getattr(_get_local_paddleocr_vl15_pipeline, "_pipeline", None)
        cached_device = getattr(_get_local_paddleocr_vl15_pipeline, "_device", None)
        if pipeline is not None and cached_device == settings["device"]:
            return pipeline

        _prepare_local_paddleocr_vl15_env(settings)
        from paddlex import create_pipeline

        pipeline = create_pipeline(
            pipeline="PaddleOCR-VL-1.5",
            device=settings["device"],
        )
        _get_local_paddleocr_vl15_pipeline._touched = True
        _get_local_paddleocr_vl15_pipeline._pipeline = pipeline
        _get_local_paddleocr_vl15_pipeline._device = settings["device"]
        return pipeline


def _prepare_local_paddleocr_vl15_env(settings: dict[str, Any]) -> None:
    os.environ[PADDLEOCR_VL15_CACHE_HOME_ENV] = str(settings["cache_home"])
    os.environ[PADDLEOCR_VL15_MODEL_SOURCE_ENV] = str(settings["model_source"])
    os.environ[PADDLEOCR_VL15_DISABLE_MODEL_SOURCE_CHECK_ENV] = (
        "true" if settings["disable_model_source_check"] else "false"
    )
    os.environ[PADDLEOCR_VL15_DISABLE_DEVICE_FALLBACK_ENV] = "true"
    import paddle

    _ensure_paddle_dynamic_mode(paddle)
    paddle.device.set_device(settings["device"])
    _patch_paddle_bfloat16_support(settings["device"])


def _get_paddleocr_vl15_runtime_settings() -> dict[str, Any]:
    raw_cache_home = Path(
        os.getenv(PADDLEOCR_VL15_CACHE_HOME_ENV, "").strip() or PROJECT_ROOT / "data" / "cache" / "paddlex"
    )
    raw_cache_home.mkdir(parents=True, exist_ok=True)
    cache_home = Path(_normalize_windows_short_path(raw_cache_home))
    cache_home.mkdir(parents=True, exist_ok=True)
    return {
        "device": os.getenv(PADDLEOCR_VL15_DEVICE_ENV, "").strip() or _detect_local_paddleocr_vl15_device(),
        "cache_home": cache_home,
        "model_source": os.getenv(PADDLEOCR_VL15_MODEL_SOURCE_ENV, "modelscope").strip() or "modelscope",
        "disable_model_source_check": _parse_bool_env(
            os.getenv(PADDLEOCR_VL15_DISABLE_MODEL_SOURCE_CHECK_ENV),
            default=True,
        ),
    }


def _detect_local_paddleocr_vl15_device() -> str:
    import paddle

    if paddle.device.is_compiled_with_cuda():
        try:
            if int(paddle.device.cuda.device_count()) > 0:
                return "gpu:0"
        except Exception:
            pass
    raise RuntimeError("当前未检测到可用 Paddle GPU 设备，VL1.5 本地模式已按 GPU 版配置，无法回退到 CPU。")


def _parse_bool_env(value: str | None, *, default: bool) -> bool:
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _normalize_windows_short_path(path: Path) -> str:
    resolved = str(path.resolve())
    if os.name != "nt":
        return resolved
    try:
        from ctypes import create_unicode_buffer, windll
    except Exception:
        return resolved

    buffer = create_unicode_buffer(4096)
    result = windll.kernel32.GetShortPathNameW(resolved, buffer, len(buffer))
    if result > 0 and buffer.value:
        return buffer.value
    return resolved


def _ensure_paddle_dynamic_mode(paddle_module: Any) -> None:
    if paddle_module.in_dynamic_mode():
        return
    paddle_module.disable_static()


def _patch_paddle_bfloat16_support(device: str) -> None:
    if getattr(_patch_paddle_bfloat16_support, "_patched_device", None) == device:
        return

    import paddle

    original = getattr(paddle.amp, "is_bfloat16_supported", None)
    if original is None:
        return

    def patched(place: Any = None) -> bool:
        if place is not None:
            return original(place)
        try:
            return original()
        except TypeError:
            pass
        except Exception as exc:
            if "Place(undefined:0)" not in str(exc):
                raise

        device_type, device_ids = _parse_vl15_device(device)
        if device_type == "gpu":
            return original(paddle.CUDAPlace(device_ids[0] if device_ids else 0))
        if device_type == "cpu":
            return original(paddle.CPUPlace())
        return False

    paddle.amp.is_bfloat16_supported = patched
    _patch_paddle_bfloat16_support._patched_device = device


def _parse_vl15_device(device: str) -> tuple[str, list[int]]:
    raw = (device or "").strip().lower()
    if raw.startswith("gpu"):
        if ":" in raw:
            try:
                return "gpu", [int(raw.split(":", 1)[1])]
            except ValueError:
                return "gpu", [0]
        return "gpu", [0]
    if raw.startswith("cpu"):
        return "cpu", []
    return raw, []

def _safe_pdf_page_count(data: bytes) -> int | None:
    try:
        return _get_pdf_page_count(data)
    except Exception:
        return None


def _slice_pdf_bytes(pdf_bytes: bytes, *, max_pages: int) -> tuple[bytes, int | None, int | None]:
    total_page_count = _safe_pdf_page_count(pdf_bytes)
    if total_page_count is None or total_page_count <= max_pages:
        return pdf_bytes, total_page_count, total_page_count
    try:
        import pymupdf
    except Exception:
        import fitz as pymupdf

    source = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        target = pymupdf.open()
        try:
            target.insert_pdf(source, from_page=0, to_page=max_pages - 1)
            return target.tobytes(garbage=4, deflate=True), total_page_count, max_pages
        finally:
            target.close()
    finally:
        source.close()


def _build_vl15_assets_root(source_bytes: bytes, filename: str) -> Path:
    payload = hashlib.sha1(source_bytes).hexdigest()
    stem = Path(filename).stem or "document"
    safe_stem = re.sub(r"[^0-9A-Za-z._-]+", "_", stem).strip("_") or "document"
    return _get_pdf_ocr_checkpoint_root() / "vl15_assets" / payload[:2] / f"{payload}_{safe_stem}"
def _document_from_ocr_result(result: Any, warnings_prefix: list[str] | None = None) -> ParsedDocument:
    pages = [
        ParsedPage(
            page_number=page.page_number,
            width=float(page.width or 0.0),
            height=float(page.height or 0.0),
            text=page.text,
            markdown=page.markdown,
            blocks=[
                ParsedBlock(
                    page_number=block.page_number,
                    block_id=block.block_id,
                    text=block.text,
                    bbox=[float(value) for value in block.bbox],
                    score=block.score,
                    block_type=block.block_type,
                    latex=block.latex,
                )
                for block in page.blocks
            ],
        )
        for page in result.pages
    ]
    warnings = [*(warnings_prefix or []), *result.warnings]
    for page in result.pages:
        warnings.extend([warning for warning in page.warnings if warning.strip()])
    return ParsedDocument(
        text=result.text,
        markdown=result.markdown,
        provider=result.provider,
        used_ocr=result.used_ocr,
        pages=pages,
        tables=[],
        warnings=warnings,
    )


def _parse_pdf(
    data: bytes,
    filename: str,
    progress_callback: ParseProgressCallback | None = None,
    cache_namespace: str | None = None,
) -> ParsedDocument:
    return _document_from_ocr_result(
        run_pdf_ocr_pipeline(
            data,
            filename,
            options=OCRPipelineOptions(
                force_ocr=True,
                render_dpi=320,
                cache_namespace=cache_namespace,
                trim_margins=True,
                remove_repeated_lines=True,
                watermark_detection=True,
            ),
            progress_callback=_ocr_progress_callback(progress_callback),
        ),
        warnings_prefix=["parse_preset=v3"],
    )


def _ocr_progress_callback(
    progress_callback: ParseProgressCallback | None,
    *,
    stage: str = "ocr",
    start: int = 15,
    end: int = 70,
):
    if progress_callback is None:
        return None

    def on_page(done_pages: int, total_pages: int, detail: dict[str, object] | None = None) -> None:
        payload = dict(detail or {})
        payload.setdefault("done_pages", done_pages)
        payload.setdefault("total_pages", total_pages)
        progress_callback(
            stage,
            min(end, start + int((done_pages / max(1, total_pages)) * max(1, end - start))),
            payload,
        )

    return on_page


def _parse_docx(data: bytes) -> ParsedDocument:
    try:
        from docx import Document

        document = Document(BytesIO(data))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        return ParsedDocument(text=text, markdown=text, provider="docx_text")
    except Exception as exc:
        warning = f"[DOCX parse failed: {exc}]"
        return ParsedDocument(text=warning, markdown=warning, provider="docx_text", warnings=[warning])


def _parse_image(data: bytes, suffix: str, *, options: DocumentParseOptions | None = None) -> ParsedDocument:
    with PADDLE_OCR_LOCK:
        parsed = _parse_with_pp_structure(data, suffix or ".png")
        if parsed is not None:
            return parsed
        return ParsedDocument(
            text="",
            markdown="",
            provider="unavailable",
            used_ocr=False,
            warnings=["PP-StructureV3 is unavailable for image OCR."],
        )


def _parse_pdf_with_pp_structure(
    data: bytes,
    filename: str,
    options: OCRPipelineOptions,
    progress_callback: ParseProgressCallback | None = None,
) -> ParsedDocument | None:
    try:
        total_page_count = _get_pdf_page_count(data)
    except Exception as exc:
        return ParsedDocument(
            text="",
            markdown="",
            provider="pp_structure_v3",
            used_ocr=True,
            warnings=[f"PP-StructureV3 PDF render failed: {exc}"],
        )

    target_page_count = min(total_page_count, options.max_pages) if options.max_pages is not None else total_page_count
    if target_page_count <= 0:
        return ParsedDocument(
            text="",
            markdown="",
            provider="pp_structure_v3",
            used_ocr=True,
            warnings=["PP-StructureV3 PDF render returned no pages."],
        )

    page_docs: list[ParsedDocument] = []
    warnings: list[str] = []
    page_batch_size = _get_pdf_page_batch_size(options)
    checkpoint_dir = _get_pdf_layout_checkpoint_dir(data, options)
    cached_pages_by_number = {
        page_number: cached_doc
        for page_number in range(1, target_page_count + 1)
        if (cached_doc := _load_pdf_layout_checkpoint_page(checkpoint_dir, page_number)) is not None
    }
    total_chunks = _count_pdf_chunks(target_page_count, page_batch_size)
    resumed_pages = len(cached_pages_by_number)
    resumed_chunks = _count_completed_chunks(cached_pages_by_number.keys(), page_batch_size)
    resume_start_page = _first_missing_page(cached_pages_by_number, target_page_count)
    pending_start_page: int | None = None
    for page_number in range(1, target_page_count + 1):
        cached_doc = cached_pages_by_number.get(page_number)
        if cached_doc is not None:
            if pending_start_page is not None:
                _process_pdf_layout_page_range(
                    data,
                    options,
                    page_docs=page_docs,
                    warnings=warnings,
                    progress_callback=progress_callback,
                    target_page_count=target_page_count,
                    page_batch_size=page_batch_size,
                    checkpoint_dir=checkpoint_dir,
                    total_chunks=total_chunks,
                    resumed_pages=resumed_pages,
                    resumed_chunks=resumed_chunks,
                    resume_start_page=resume_start_page,
                    start_page=pending_start_page,
                    end_page=page_number - 1,
                )
                pending_start_page = None
            page_docs.append(cached_doc)
            if progress_callback is not None:
                progress_callback(
                    "layout_analysis",
                    min(70, 15 + int((page_number / max(1, target_page_count)) * 55)),
                    {
                        "done_pages": len(page_docs),
                        "total_pages": target_page_count,
                        "page_batch_size": page_batch_size,
                        "chunk_count": total_chunks,
                        "completed_chunk_count": _count_fully_completed_chunks(
                            len(page_docs),
                            target_page_count,
                            page_batch_size,
                        ),
                        "resumed_pages": resumed_pages,
                        "resumed_chunk_count": resumed_chunks,
                        "resume_start_page": resume_start_page,
                    },
                )
            continue

        if pending_start_page is None:
            pending_start_page = page_number

    if pending_start_page is not None:
        _process_pdf_layout_page_range(
            data,
            options,
            page_docs=page_docs,
            warnings=warnings,
            progress_callback=progress_callback,
            target_page_count=target_page_count,
            page_batch_size=page_batch_size,
            checkpoint_dir=checkpoint_dir,
            total_chunks=total_chunks,
            resumed_pages=resumed_pages,
            resumed_chunks=resumed_chunks,
            resume_start_page=resume_start_page,
            start_page=pending_start_page,
            end_page=target_page_count,
        )

    if not page_docs:
        return ParsedDocument(
            text="",
            markdown="",
            provider="pp_structure_v3",
            used_ocr=True,
            warnings=warnings or ["PP-StructureV3 returned no page documents."],
        )

    pages_out = [page for doc in page_docs for page in doc.pages]
    tables = [table for doc in page_docs for table in doc.tables]
    text = "\n\n".join(doc.text for doc in page_docs if doc.text.strip())
    markdown = "\n\n".join(doc.markdown for doc in page_docs if doc.markdown.strip())
    markdown_image_roots = [
        root
        for doc in page_docs
        for root in (getattr(doc, "markdown_image_roots", []) or [])
        if str(root).strip()
    ]
    if options.max_pages is not None and total_page_count > target_page_count:
        warnings.append(f"Layout analysis limited to first {target_page_count} of {total_page_count} pages.")
    return ParsedDocument(
        text=text,
        markdown=markdown or text,
        provider="pp_structure_v3",
        used_ocr=True,
        pages=pages_out,
        tables=tables,
        warnings=warnings,
        markdown_image_roots=markdown_image_roots,
    )


def _process_pdf_layout_page_range(
    data: bytes,
    options: OCRPipelineOptions,
    *,
    page_docs: list[ParsedDocument],
    warnings: list[str],
    progress_callback: ParseProgressCallback | None,
    target_page_count: int,
    page_batch_size: int,
    checkpoint_dir: Path,
    total_chunks: int,
    resumed_pages: int,
    resumed_chunks: int,
    resume_start_page: int | None,
    start_page: int,
    end_page: int,
) -> None:
    for batch_start in range(start_page - 1, end_page, page_batch_size):
        batch_end = min(batch_start + page_batch_size, end_page)
        pages, render_warnings = _render_pdf_page_range_resilient(
            data,
            options.render_dpi,
            batch_start,
            batch_end,
        )
        warnings.extend(render_warnings)

        for page in pages:
            page_number = int(page.get("page_number") or 0)
            if progress_callback is not None:
                progress_callback(
                    "layout_analysis",
                    min(70, 15 + int((page_number / max(1, target_page_count)) * 55)),
                    {
                        "done_pages": len(page_docs) + 1,
                        "total_pages": target_page_count,
                        "page_batch_size": page_batch_size,
                        "chunk_count": total_chunks,
                        "completed_chunk_count": _count_fully_completed_chunks(
                            len(page_docs) + 1,
                            target_page_count,
                            page_batch_size,
                        ),
                        "current_chunk_index": (batch_start // page_batch_size) + 1,
                        "current_chunk_page_from": batch_start + 1,
                        "current_chunk_page_to": batch_end,
                        "resumed_pages": resumed_pages,
                        "resumed_chunk_count": resumed_chunks,
                        "resume_start_page": resume_start_page,
                    },
                )
            page_doc = _parse_with_pp_structure(
                page["image_bytes"],
                ".png",
                enable_formula_recognition=options.enable_formula_recognition,
            )
            if page_doc is None:
                warnings.append(f"PP-StructureV3 unavailable on page {page_number}.")
                continue
            if page_doc.warnings:
                warnings.extend(f"page {page_number}: {warning}" for warning in page_doc.warnings)
            _save_pdf_layout_checkpoint_assets(checkpoint_dir, page_number, page_doc)
            for parsed_page in page_doc.pages:
                parsed_page.page_number = page_number
            for table in page_doc.tables:
                table.page = page_number
            page_docs.append(page_doc)
            _save_pdf_layout_checkpoint_page(checkpoint_dir, page_number, page_doc)

        release_paddle_ocr_resources(clear_cached_engines=False, force_clear_cache=True)


def _get_pdf_layout_checkpoint_dir(data: bytes, options: OCRPipelineOptions) -> Path:
    options_payload = {
        key: value
        for key, value in asdict(options).items()
        if key not in {"max_pages", "page_chunk_size"}
    }
    payload = json.dumps(
        {
            "provider": "pp_structure_v3",
            "sha256": hashlib.sha256(data).hexdigest(),
            "options": options_payload,
            "cache_namespace": options.cache_namespace,
        },
        ensure_ascii=False,
        sort_keys=True,
    )
    digest = hashlib.sha1(payload.encode("utf-8")).hexdigest()
    checkpoint_dir = _get_pdf_ocr_checkpoint_root() / "layout" / digest[:2] / digest
    checkpoint_dir.mkdir(parents=True, exist_ok=True)
    (checkpoint_dir / "source.sha256").write_text(hashlib.sha256(data).hexdigest(), encoding="utf-8")
    if options.cache_namespace:
        namespace_path = checkpoint_dir / CHECKPOINT_NAMESPACE_FILENAME
        if not namespace_path.exists():
            namespace_path.write_text(options.cache_namespace, encoding="utf-8")
    return checkpoint_dir


def _load_pdf_layout_checkpoint_page(checkpoint_dir: Path, page_number: int) -> ParsedDocument | None:
    path = checkpoint_dir / f"page_{page_number:05d}.json"
    if not path.exists():
        return None
    parsed = deserialize_parsed_document(path.read_text(encoding="utf-8"))
    if parsed is None:
        return None
    if _document_has_image_refs_but_no_assets(parsed):
        return None
    return parsed


def _save_pdf_layout_checkpoint_page(checkpoint_dir: Path, page_number: int, document: ParsedDocument) -> None:
    if any((warning or "").startswith("PP-StructureV3 parse failed:") for warning in document.warnings):
        return
    path = checkpoint_dir / f"page_{page_number:05d}.json"
    tmp_path = path.with_suffix(".json.tmp")
    tmp_path.write_text(serialize_parsed_document(document), encoding="utf-8")
    tmp_path.replace(path)


def _save_pdf_layout_checkpoint_assets(checkpoint_dir: Path, page_number: int, document: ParsedDocument) -> None:
    markdown_images = dict(getattr(document, "markdown_images", {}) or {})
    if not markdown_images:
        return
    assets_root = checkpoint_dir / f"page_{page_number:05d}_assets"
    roots = _persist_markdown_images(markdown_images, assets_root)
    document.markdown_image_roots.extend(root for root in roots if root not in document.markdown_image_roots)
    document.markdown_images = {}


def _parse_text_document(data: bytes) -> ParsedDocument:
    text = _parse_text(data)
    return ParsedDocument(text=text, markdown=text, provider="plain_text")


def _parse_text(data: bytes) -> str:
    for encoding in ("utf-8", "gb18030", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_selectable_pdf_text(data: bytes) -> ParsedDocument | None:
    try:
        import fitz

        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        return ParsedDocument(
            text="",
            markdown="",
            provider="pymupdf_text",
            warnings=[f"PyMuPDF text extraction failed: {exc}"],
        )

    pages: list[ParsedPage] = []
    page_texts: list[str] = []
    try:
        for index, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            raw_blocks = page.get_text("blocks") or []
            blocks: list[ParsedBlock] = []
            for block_index, raw_block in enumerate(raw_blocks, start=1):
                if not isinstance(raw_block, (list, tuple)) or len(raw_block) < 5:
                    continue
                block_text = str(raw_block[4] or "").strip()
                if not block_text:
                    continue
                blocks.append(
                    ParsedBlock(
                        page_number=index,
                        block_id=f"p{index}-b{block_index}",
                        text=block_text,
                        bbox=[
                            float(raw_block[0]),
                            float(raw_block[1]),
                            float(raw_block[2]),
                            float(raw_block[3]),
                        ],
                    )
                )
            if not blocks and text:
                blocks = _synthetic_blocks_from_text(text, index)
            pages.append(
                ParsedPage(
                    page_number=index,
                    width=float(page.rect.width),
                    height=float(page.rect.height),
                    text=text,
                    markdown=text,
                    blocks=blocks,
                )
            )
            if text:
                page_texts.append(text)
    finally:
        doc.close()

    merged = "\n\n".join(page_texts)
    return ParsedDocument(
        text=merged,
        markdown=merged,
        provider="pymupdf_text",
        used_ocr=False,
        pages=pages,
    )


def _document_has_meaningful_text(document: ParsedDocument | None) -> bool:
    if document is None:
        return False
    return len(TEXT_CHAR_PATTERN.findall(document.text or "")) >= PDF_TEXT_CHAR_THRESHOLD


def _document_quality_score(document: ParsedDocument) -> float:
    text = (document.markdown or document.text or "").strip()
    if not text:
        return 0.0
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    text_chars = len(TEXT_CHAR_PATTERN.findall(text))
    question_hits = len(re.findall(r"(?m)^\s*\d{1,3}\s*[\.、．)]", text))
    option_hits = len(re.findall(r"(?m)^\s*[A-H][\.\、．)]", text))
    answer_hits = len(re.findall(r"(?:答案|参考答案|正确答案)\s*[:：]", text))
    analysis_hits = len(re.findall(r"(?:解析|答案解析|【解析】)\s*[:：]", text))
    short_line_ratio = (
        sum(1 for line in lines if len(TEXT_CHAR_PATTERN.findall(line)) <= 3) / max(1, len(lines))
    )
    score = 0.0
    score += min(3.0, text_chars / 160)
    score += min(3.0, question_hits * 0.9)
    score += min(2.0, option_hits * 0.25)
    score += min(1.0, answer_hits * 0.5)
    score += min(1.0, analysis_hits * 0.5)
    score -= min(2.0, short_line_ratio * 2.0)
    return max(0.0, round(score, 2))


def _document_has_unresolved_formula_images(document: ParsedDocument | None) -> bool:
    if document is None:
        return False
    payload_parts = [
        str(part or "")
        for part in (document.raw_markdown, document.markdown, document.raw_text, document.text)
        if str(part or "").strip()
    ]
    payload_parts.extend(
        str(getattr(page, "markdown", "") or getattr(page, "text", "") or "")
        for page in (document.pages or [])
    )
    payload = "\n".join(part for part in payload_parts if part.strip())
    return bool(FORMULA_IMAGE_REF_PATTERN.search(payload))


def _parse_with_pp_structure(
    data: bytes,
    suffix: str,
    *,
    enable_formula_recognition: bool = False,
) -> ParsedDocument | None:
    pipeline = _get_pp_structure_pipeline()
    if pipeline is None:
        return None

    tmp_path: str | None = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        with os.fdopen(fd, "wb") as tmp:
            tmp.write(data)
            tmp.flush()
        results = _predict_pp_structure(
            pipeline,
            tmp_path,
            enable_formula_recognition=enable_formula_recognition,
        )
    except Exception as exc:
        return ParsedDocument(
            text="",
            markdown="",
            provider="pp_structure_v3",
            used_ocr=True,
            warnings=[f"PP-StructureV3 parse failed: {exc}"],
        )
    finally:
        if tmp_path:
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    return _document_from_pp_results(results, provider="pp_structure_v3")


def _predict_pp_structure(pipeline: Any, input_path: str, *, enable_formula_recognition: bool):
    with PADDLE_OCR_LOCK:
        return list(
            pipeline.predict(
                input=input_path,
                use_textline_orientation=True,
                use_table_recognition=True,
                use_formula_recognition=enable_formula_recognition,
            )
        )


def _document_from_pp_results(results: list[Any], provider: str) -> ParsedDocument:
    pages: list[ParsedPage] = []
    tables: list[ParsedTable] = []
    warnings: list[str] = []
    text_parts: list[str] = []
    markdown_parts: list[str] = []
    markdown_images: dict[str, Any] = {}

    for page_index, result in enumerate(results, start=1):
        text = _extract_pp_text(result).strip()
        markdown, page_markdown_images = _extract_pp_markdown_bundle(result)
        markdown = markdown.strip()
        if page_markdown_images:
            prefixed_images, path_mapping = _prefix_markdown_image_paths(page_markdown_images, page_number=page_index)
            markdown_images.update(prefixed_images)
            text = _replace_markdown_image_paths(text, path_mapping)
            markdown = _replace_markdown_image_paths(markdown, path_mapping)
        page_tables = _extract_pp_tables(result, page_index)
        tables.extend(page_tables)
        text = _dedupe_table_markup(text, page_tables)
        markdown = _dedupe_table_markup(markdown, page_tables)
        pages.append(
            ParsedPage(
                page_number=page_index,
                text=text,
                markdown=markdown or text,
                blocks=_synthetic_blocks_from_text(text, page_index),
            )
        )
        if text:
            text_parts.append(text)
        if markdown:
            markdown_parts.append(markdown)
        elif text:
            markdown_parts.append(text)

    merged_text = "\n\n".join(part for part in text_parts if part)
    merged_markdown = "\n\n".join(part for part in markdown_parts if part)
    if not merged_text and not merged_markdown:
        warnings.append("PP-StructureV3 returned no textual content.")

    return ParsedDocument(
        text=merged_text or _markdown_to_text(merged_markdown),
        markdown=merged_markdown or merged_text,
        provider=provider,
        used_ocr=True,
        pages=pages,
        tables=tables,
        warnings=warnings,
        markdown_image_roots=[],
        markdown_images=markdown_images,
    )


def _extract_pp_text(result: Any) -> str:
    for candidate in (
        getattr(result, "text", None),
        getattr(result, "markdown", None),
        getattr(result, "rec_texts", None),
        getattr(result, "texts", None),
    ):
        normalized = _normalize_text_candidate(candidate)
        if normalized.strip():
            return _markdown_to_text(normalized)

    payload = _result_to_dict(result)
    for key in ("text", "markdown", "rec_texts", "texts"):
        normalized = _normalize_text_candidate(payload.get(key))
        if normalized.strip():
            return _markdown_to_text(normalized)

    outputs = payload.get("parsing_res_list") or payload.get("layout_parsing_result") or payload.get("results")
    return _collect_text_from_blocks(outputs)


def _extract_pp_markdown(result: Any) -> str:
    markdown, _ = _extract_pp_markdown_bundle(result)
    return markdown


def _extract_pp_markdown_bundle(result: Any) -> tuple[str, dict[str, Any]]:
    markdown_value = getattr(result, "markdown", None)
    if isinstance(markdown_value, dict):
        markdown = _extract_markdown_payload(markdown_value)
        markdown_images = _extract_markdown_image_payload(markdown_value.get("markdown_images"))
        if markdown.strip() or markdown_images:
            return markdown, markdown_images
    else:
        markdown = _extract_markdown_payload(markdown_value)
        if markdown.strip():
            return markdown, {}

    payload = _result_to_dict(result)
    for key in ("markdown", "md"):
        candidate = payload.get(key)
        if isinstance(candidate, dict):
            markdown = _extract_markdown_payload(candidate)
            markdown_images = _extract_markdown_image_payload(candidate.get("markdown_images"))
            if markdown.strip() or markdown_images:
                return markdown, markdown_images
        else:
            markdown = _extract_markdown_payload(candidate)
            if markdown.strip():
                return markdown, {}

    outputs = payload.get("parsing_res_list") or payload.get("layout_parsing_result") or payload.get("results")
    return _collect_markdown_from_blocks(outputs), {}


def _extract_pp_tables(result: Any, page_number: int) -> list[ParsedTable]:
    payload = _result_to_dict(result)
    tables: list[ParsedTable] = []
    table_res_list = payload.get("table_res_list")
    if isinstance(table_res_list, list):
        for item in table_res_list:
            if not isinstance(item, dict):
                continue
            html = _normalize_text_candidate(item.get("pred_html") or item.get("html") or item.get("table_html"))
            markdown = _normalize_text_candidate(item.get("pred_markdown") or item.get("markdown") or item.get("table_markdown"))
            if html or markdown:
                tables.append(ParsedTable(page=page_number, markdown=markdown, html=html))

    outputs = payload.get("parsing_res_list") or payload.get("layout_parsing_result") or payload.get("results")
    if not isinstance(outputs, list):
        return tables

    for item in outputs:
        if not isinstance(item, dict):
            continue
        block_type = str(item.get("block_label") or item.get("type") or "").lower()
        if "table" not in block_type:
            continue
        markdown = _normalize_text_candidate(
            item.get("markdown") or item.get("table_markdown") or item.get("pred_markdown")
        )
        html = _normalize_text_candidate(item.get("html") or item.get("table_html") or item.get("pred_html"))
        if markdown or html:
            signature = (markdown.strip(), html.strip())
            if not any((existing.markdown.strip(), existing.html.strip()) == signature for existing in tables):
                tables.append(ParsedTable(page=page_number, markdown=markdown, html=html))
    return tables


def _result_to_dict(result: Any) -> dict[str, Any]:
    if isinstance(result, dict):
        return result
    if hasattr(result, "__dict__"):
        return dict(vars(result))
    return {}


def _collect_text_from_blocks(outputs: Any) -> str:
    if not isinstance(outputs, list):
        return ""
    lines: list[str] = []
    for item in outputs:
        if not isinstance(item, dict):
            continue
        for key in ("text", "markdown", "content", "result", "latex"):
            value = _normalize_text_candidate(item.get(key))
            if value.strip():
                lines.append(_markdown_to_text(value))
                break
    return "\n".join(line for line in lines if line).strip()


def _collect_markdown_from_blocks(outputs: Any) -> str:
    if not isinstance(outputs, list):
        return ""
    blocks: list[str] = []
    for item in outputs:
        if not isinstance(item, dict):
            continue
        value = _normalize_text_candidate(
            item.get("markdown") or item.get("table_markdown") or item.get("content") or item.get("text")
        )
        if value.strip():
            blocks.append(value)
    return "\n\n".join(block for block in blocks if block).strip()


def _normalize_text_candidate(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        preferred = []
        for key in ("markdown_texts", "text", "block_content", "content"):
            item = value.get(key)
            text = _normalize_text_candidate(item)
            if text:
                preferred.append(text)
        if preferred:
            return "\n".join(preferred).strip()
        return str(value).strip()
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (list, tuple)):
        parts = [_normalize_text_candidate(item) for item in value]
        return "\n".join(part for part in parts if part).strip()
    return str(value).strip()


def _synthetic_blocks_from_text(text: str, page_number: int) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    for index, line in enumerate(text.splitlines(), start=1):
        normalized = line.strip()
        if not normalized:
            continue
        blocks.append(
            ParsedBlock(
                page_number=page_number,
                block_id=f"p{page_number}-synthetic-{index}",
                text=normalized,
            )
        )
    return blocks


def _markdown_to_text(markdown: str) -> str:
    if not markdown:
        return ""
    text = markdown
    text = re.sub(r"```.*?```", "", text, flags=re.S)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^[>#\-\*\s|]+", "", text, flags=re.M)
    text = re.sub(r"\|", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_markdown_payload(value: Any) -> str:
    if isinstance(value, dict):
        return _normalize_text_candidate(value.get("markdown_texts") or value.get("text") or value)
    return _normalize_text_candidate(value)


def _extract_markdown_image_payload(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    return {
        str(path).replace("\\", "/"): image
        for path, image in value.items()
        if str(path).strip() and image is not None
    }


def _prefix_markdown_image_paths(markdown_images: dict[str, Any], *, page_number: int) -> tuple[dict[str, Any], dict[str, str]]:
    prefixed: dict[str, Any] = {}
    path_mapping: dict[str, str] = {}
    for original_path, image in markdown_images.items():
        normalized_path = str(original_path).replace("\\", "/").strip()
        if not normalized_path:
            continue
        prefixed_path = f"imgs/page_{page_number:04d}/{Path(normalized_path).name}"
        prefixed[prefixed_path] = image
        path_mapping[normalized_path] = prefixed_path
    return prefixed, path_mapping


def _replace_markdown_image_paths(content: str, path_mapping: dict[str, str]) -> str:
    if not content or not path_mapping:
        return content
    updated = content
    for original_path, next_path in sorted(path_mapping.items(), key=lambda item: len(item[0]), reverse=True):
        updated = updated.replace(original_path, next_path)
    return updated


def _strip_pdf_embedded_image_content(document: ParsedDocument) -> ParsedDocument:
    document.text = _strip_pdf_image_markup(document.text)
    document.markdown = _strip_pdf_image_markup(document.markdown)
    document.raw_text = _strip_pdf_image_markup(document.raw_text)
    document.raw_markdown = _strip_pdf_image_markup(document.raw_markdown)
    document.markdown_images = {}
    document.markdown_image_roots = []
    for page in document.pages:
        page.text = _strip_pdf_image_markup(page.text)
        page.markdown = _strip_pdf_image_markup(page.markdown)
        for block in page.blocks:
            block.text = _strip_pdf_image_markup(block.text)
    for table in document.tables:
        table.markdown = _strip_pdf_image_markup(table.markdown)
        table.html = PDF_IMAGE_HTML_PATTERN.sub("", table.html or "").strip()
    return document


def _strip_pdf_image_markup(content: str) -> str:
    if not content:
        return content
    updated = PDF_IMAGE_MARKDOWN_PATTERN.sub("", content)
    updated = PDF_IMAGE_HTML_PATTERN.sub("", updated)
    updated = PDF_IMAGE_PATH_PATTERN.sub("", updated)
    updated = re.sub(r"[ \t]+\n", "\n", updated)
    updated = re.sub(r"\n{3,}", "\n\n", updated)
    return updated.strip()


def _persist_markdown_images(markdown_images: dict[str, Any], output_root: Path) -> list[str]:
    if not markdown_images:
        return []
    output_root.mkdir(parents=True, exist_ok=True)
    for relative_path, image in markdown_images.items():
        target_path = output_root / Path(relative_path.replace("\\", "/"))
        target_path.parent.mkdir(parents=True, exist_ok=True)
        if hasattr(image, "save"):
            image.save(target_path)
            continue
        if isinstance(image, (bytes, bytearray)):
            target_path.write_bytes(bytes(image))
            continue
        if isinstance(image, str):
            raw = image.strip()
            if not raw:
                continue
            if raw.startswith("data:") and "," in raw:
                raw = raw.split(",", 1)[1]
            try:
                target_path.write_bytes(base64.b64decode(raw))
            except Exception:
                continue
    return [str(output_root)]


def _document_has_image_refs_but_no_assets(document: ParsedDocument | None) -> bool:
    if document is None:
        return False
    payload = "\n".join(
        str(part or "")
        for part in (document.raw_markdown, document.markdown, document.raw_text, document.text)
        if str(part or "").strip()
    )
    return ('src="imgs/' in payload or "src='imgs/" in payload) and not (document.markdown_image_roots or [])


def _dedupe_table_markup(content: str, tables: list[ParsedTable]) -> str:
    if not content:
        return content
    normalized = content
    for table in tables:
        html = (table.html or "").strip()
        if html:
            normalized = normalized.replace(html, "").replace(f"<div style=\"text-align: center;\">{html}</div>", "")
        markdown = (table.markdown or "").strip()
        if markdown:
            normalized = normalized.replace(markdown, "")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _get_pp_structure_pipeline():
    try:
        _patch_paddleocr_common_args()
        from paddleocr import PPStructureV3
    except Exception:
        return None

    if getattr(_get_pp_structure_pipeline, "_pipeline", None) is None:
        with PADDLE_OCR_LOCK:
            if getattr(_get_pp_structure_pipeline, "_pipeline", None) is None:
                _get_pp_structure_pipeline._touched = True
                ocr_settings = _get_paddle_ocr_settings()
                init_kwargs = {
                    "device": str(ocr_settings["device"]),
                    "enable_hpi": False,
                    "use_tensorrt": False,
                    "enable_mkldnn": False,
                    "cpu_threads": 4,
                    "text_rec_score_thresh": 0.0,
                }
                if ocr_settings["text_detection_model_dir"]:
                    init_kwargs["text_detection_model_dir"] = ocr_settings["text_detection_model_dir"]
                else:
                    init_kwargs["text_detection_model_name"] = ocr_settings["text_detection_model_name"]
                if ocr_settings["text_recognition_model_dir"]:
                    init_kwargs["text_recognition_model_dir"] = ocr_settings["text_recognition_model_dir"]
                else:
                    init_kwargs["text_recognition_model_name"] = ocr_settings["text_recognition_model_name"]
                pipeline = PPStructureV3(**init_kwargs)
                _get_pp_structure_pipeline._pipeline = pipeline
    return _get_pp_structure_pipeline._pipeline


def release_paddle_parser_resources() -> None:
    with PADDLE_OCR_LOCK:
        had_pp_structure = (
            getattr(_get_pp_structure_pipeline, "_pipeline", None) is not None
            or getattr(_get_pp_structure_pipeline, "_touched", False)
        )
        local_vl15_pipeline = getattr(_get_local_paddleocr_vl15_pipeline, "_pipeline", None)
        had_local_vl15 = (
            local_vl15_pipeline is not None
            or getattr(_get_local_paddleocr_vl15_pipeline, "_touched", False)
        )
        if local_vl15_pipeline is not None and hasattr(local_vl15_pipeline, "close"):
            try:
                local_vl15_pipeline.close()
            except Exception:
                pass
        _get_local_paddleocr_vl15_pipeline._pipeline = None
        _get_local_paddleocr_vl15_pipeline._device = None
        _get_local_paddleocr_vl15_pipeline._touched = False
        _get_pp_structure_pipeline._pipeline = None
        _get_pp_structure_pipeline._touched = False
        release_paddle_ocr_resources(force_clear_cache=(had_pp_structure or had_local_vl15))


def _patch_paddleocr_common_args() -> None:
    if getattr(_patch_paddleocr_common_args, "_patched", False):
        return

    try:
        from paddlex.inference import PaddlePredictorOption
        from paddlex.utils.device import get_default_device, parse_device
        from paddleocr import _common_args as common_args
    except Exception:
        return

    def prepare_common_init_args(model_name, common_args_dict):
        device = common_args_dict["device"]
        if device is None:
            device = get_default_device()
        device_type, device_ids = parse_device(device)
        device_id = device_ids[0] if device_ids is not None else None

        init_kwargs = {
            "use_hpip": common_args_dict["enable_hpi"],
            "hpi_config": {"device_type": device_type, "device_id": device_id},
        }
        pp_option = PaddlePredictorOption(device_type=device_type, device_id=device_id)
        if model_name:
            pp_option.setdefault_by_model_name(model_name)

        if device_type == "gpu":
            if common_args_dict["use_pptrt"]:
                if common_args_dict["pptrt_precision"] == "fp32":
                    pp_option.run_mode = "trt_fp32"
                else:
                    pp_option.run_mode = "trt_fp16"
            else:
                pp_option.run_mode = "paddle"
        elif device_type == "cpu":
            if common_args_dict["enable_mkldnn"]:
                pp_option.run_mode = "mkldnn"
                pp_option.mkldnn_cache_capacity = common_args_dict["mkldnn_cache_capacity"]
            else:
                pp_option.run_mode = "paddle"
            pp_option.cpu_threads = common_args_dict["cpu_threads"]
        else:
            pp_option.run_mode = "paddle"

        init_kwargs["pp_option"] = pp_option
        return init_kwargs

    common_args.prepare_common_init_args = prepare_common_init_args
    _patch_paddleocr_common_args._patched = True
